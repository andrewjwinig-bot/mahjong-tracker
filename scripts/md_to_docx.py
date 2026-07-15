#!/usr/bin/env python3
"""Convert the CPA job description markdown into a formatted Word document."""
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = sys.argv[1]
OUT = sys.argv[2]

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)  # deep navy

doc = Document()

# Base font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def add_runs(paragraph, text):
    """Render **bold** and *italic* inline markup into runs."""
    for part in re.split(r"(\*\*.*?\*\*|\*.*?\*)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("*") and part.endswith("*"):
            r = paragraph.add_run(part[1:-1])
            r.italic = True
        else:
            paragraph.add_run(part)


with open(SRC) as f:
    lines = f.read().splitlines()

i = 0
while i < len(lines):
    line = lines[i].rstrip()

    if not line:
        i += 1
        continue

    # Horizontal rule -> subtle spacer, skip the divider itself
    if line.strip() == "---":
        i += 1
        continue

    # Headings
    if line.startswith("### "):
        p = doc.add_heading(level=2)
        add_runs(p, line[4:])
        for r in p.runs:
            r.font.color.rgb = ACCENT
        i += 1
        continue
    if line.startswith("## "):
        p = doc.add_heading(level=1)
        add_runs(p, line[3:])
        for r in p.runs:
            r.font.color.rgb = ACCENT
        i += 1
        continue
    if line.startswith("# "):
        p = doc.add_heading(level=0)
        add_runs(p, line[2:])
        i += 1
        continue

    # Nested bullet (sub-item)
    m_sub = re.match(r"^ {2,}- (.*)", lines[i])
    if m_sub:
        p = doc.add_paragraph(style="List Bullet 2")
        add_runs(p, m_sub.group(1))
        i += 1
        continue

    # Bullet
    if line.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        add_runs(p, line[2:])
        i += 1
        continue

    # Plain paragraph (join wrapped lines until blank/structural)
    buf = [line]
    j = i + 1
    while j < len(lines):
        nxt = lines[j].rstrip()
        if (not nxt or nxt.strip() == "---" or nxt.startswith("#")
                or nxt.startswith("- ") or re.match(r"^ {2,}- ", lines[j])):
            break
        buf.append(nxt)
        j += 1
    p = doc.add_paragraph()
    add_runs(p, " ".join(buf))
    i = j

doc.save(OUT)
print("Wrote", OUT)
